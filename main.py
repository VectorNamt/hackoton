import docx
import os
import re
import nltk
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
from nltk.corpus import stopwords
nltk.download('punkt')
nltk.download('stopwords')

paths = str(input("Введите путь к файлу\n"))
folder = os.getcwd()
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))


for path in paths:
    doc = docx.Document(paths)
    
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)


for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text.append(cell.text)



bad_chars = [';', ':', '!', '*', '№', '$', '%', '^', '&', '?', '*', ')', '(', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9' '-', '+', '_', '+', '{', '}', '[', ']', '/', '.', ',', '<', '>']



text = '_'.join(text)

stop_words = set(stopwords.words("russian")) #создание списка на нежелательных слов

keyword = [         #список ключевых слов
    [Интеллект, искуствен, ПО, Нейрон, машин, обучен, развит, подход, алгоритм, интеллектуль, способн, сам, использ, мозг, обработк],
    [технолог,разработк,продукц,процесс,химическ,техник,металл,обработка,средств,строительнств,космическ,практическ,определен,производств,материал,управлен,научн,научно-техническ,деятельн],
    [робототехник,робот,развит,сенсорик,област,субтехнолог,сенсор,цифров,обработк,информац,взаимодейств,человек,просмотр,датчик,по,моделирован,проектирован,механическ,рынок,инновацион,промышл],
    [Средств,интернет, сет,вещ,концеп,устройств,беспроводн,индентификац,физич,уровн,прибор,компан,реальн,будущ,развит,самртфон,деньг,товар,безопасн,товар,дист],
    [5g,3g,4g,сет,связ,технолог, мобильн,скорост,частот,устройств,диапазон,передач,интернет,оператор,задержк,мудл,поколен,сотов,станц,базов,антен,волн,LTE,ГГц],
    [технолог,информац,информацион,коммуникац,протокол,передач,современ,процесс,пользовател,сервер,электрон,социальн,сообщен,ресурс,обеспечива,образован,обучен,информационно,коммуникацион,коммуникативн,сетев,беспроводн],
    [реальност,виртуальн,vr,дополнен,пользовател,реальн,ar,объект,обучен,позволя,сфер,vr/ar,погружен,применен,информац,взаимодейств,пространств,игр,reality,компьютерн,смартфон,дополнительн],
    [блокчейн,реестр,сет,технолог,распределен,транзакц,консенсус,блок,ibm,компан,безопасн,баз,ключ,алгоритм,финансов,крипт,hyperledger,децентрализова,хран,криптовалют,крипт,dlt],
    [квант, криптограф, фотон, коэффициент, поляризац, модул, коммуникац, работ, технолог, шифрован, алгоритм],
    [квант, квадратическ, сенсор, измерен, измерительн, величин, высокоточн, относительн, суперпозиц, технолог],
    [кубит, квант, квантов, суперпозиц, вычислительн, детерминирова, многокубитн, вероятн, экспоненциальн, электрон, фотон],
     
    ]

words = nltk.word_tokenize(text)
without_stop_words = [word for word in words if not word in stop_words]

without_stop_words = '_'.join(without_stop_words)

for i in bad_chars : 
    without_stop_words = without_stop_words.replace(i, ' ')

MainText = f'List of Words ={without_stop_words.split()}'

FileTheame = int(input(Введите тему программы)) #ввод темы для скана
      
look_like_proc = 0 #создание переменной соответствия
    
for i in keyword:       #получение процента соответвия
    for j in MainText:
        look_like = fuzz.ratio(keyword[FileTheame[i]], MainText[j])
        if look_like > 40:
            look_like_proc += 100/len(keyword)
            





